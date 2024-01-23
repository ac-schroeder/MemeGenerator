import unittest
import pathlib
import os
import pandas
import subprocess
import docx

# import app
from QuoteEngine import Ingestor
from QuoteEngine import QuoteModel
from QuoteEngine import TextIngestor
from QuoteEngine import CSVIngestor
from QuoteEngine import PDFIngestor
from QuoteEngine import DocxIngestor

TESTS_ROOT = (pathlib.Path(__file__).parent).resolve()
TEST_CSV_FILE = TESTS_ROOT.joinpath('DogQuotes/DogQuotesCSV.csv')
TEST_DOCX_FILE = TESTS_ROOT.joinpath('DogQuotes/DogQuotesDOCX.docx')
TEST_PDF_FILE = TESTS_ROOT.joinpath('DogQuotes/DogQuotesPDF.pdf')
TEST_TXT_FILE = TESTS_ROOT.joinpath('DogQuotes/DogQuotesTXT.txt')
TEST_JSON_FILE = TESTS_ROOT.joinpath('DogQuotes/DogQuotesJSON.json')


class TestIngestor(unittest.TestCase):

    # Ingestor Tests #

    def test_ingestor_returns_empty_quotes_for_unsupported_file(self):
        quotes = Ingestor.parse(TEST_JSON_FILE)
        self.assertFalse(quotes)

    # Text Ingestor Tests #

    def test_text_ingestor_returns_quote_objects(self):
        quotes = Ingestor.parse(TEST_TXT_FILE)
        self.assertIsInstance(quotes[0], QuoteModel)

    def test_text_ingestor_returns_all_quotes(self):
        with open(TEST_TXT_FILE, encoding='utf-8-sig') as f:
            lines = [line for line in f if line]
        quotes = Ingestor.parse(TEST_TXT_FILE)
        self.assertEqual(len(lines), len(quotes))

    def test_text_ingestor_quotes_match_content(self):
        with open(TEST_TXT_FILE, encoding='utf-8-sig') as f:
            line = f.readline()
            line = line.strip('\n\r').strip()
            text_quote_body = line.split(" - ")[0].strip('"\'')
            text_quote_author = line.split(" - ")[1]

        quotes = Ingestor.parse(TEST_TXT_FILE)
        quote_obj_body = quotes[0].body.strip('"\'')
        quote_obj_author = quotes[0].author

        self.assertEqual(quote_obj_author, text_quote_author)
        self.assertEqual(quote_obj_body, text_quote_body)

    def test_text_ingestor_raises_error_for_wrong_filetype(self):
        self.assertRaises(Exception,
                          TextIngestor.TextIngestor.parse,
                          TEST_JSON_FILE)

    # CSV Ingestor Tests #

    def test_csv_ingestor_returns_quote_objects(self):
        quotes = Ingestor.parse(TEST_CSV_FILE)
        self.assertIsInstance(quotes[0], QuoteModel)

    def test_csv_ingestor_returns_all_quotes(self):
        df = pandas.read_csv(TEST_CSV_FILE, header=0)
        quotes = Ingestor.parse(TEST_CSV_FILE)
        self.assertEqual(df.shape[0], len(quotes))

    def test_csv_ingestor_quotes_match_content(self):
        df = pandas.read_csv(TEST_CSV_FILE, header=0)
        csv_quote_body = df['body'].iloc[0].strip('"\'')
        csv_quote_author = df['author'].iloc[0]

        quotes = Ingestor.parse(TEST_CSV_FILE)
        quote_obj_body = quotes[0].body.strip('"\'')
        quote_obj_author = quotes[0].author

        self.assertEqual(quote_obj_author, csv_quote_author)
        self.assertEqual(quote_obj_body, csv_quote_body)

    def test_csv_ingestor_raises_error_for_wrong_filetype(self):
        self.assertRaises(Exception,
                          CSVIngestor.CSVIngestor.parse,
                          TEST_JSON_FILE)

    # PDF Ingestor Tests #

    def test_pdf_ingestor_returns_quote_objects(self):
        quotes = Ingestor.parse(TEST_PDF_FILE)
        self.assertIsInstance(quotes[0], QuoteModel)

    def test_pdf_ingestor_returns_all_quotes(self):
        tmp = TEST_PDF_FILE.parent.joinpath('tmp.txt')
        subprocess.run(['pdftotext', '-layout',
                        TEST_PDF_FILE, tmp])
        with open(tmp) as f:
            lines = [line for line in f
                     if len(line.strip('\n\r').strip()) > 0]
        os.remove(tmp)

        quotes = Ingestor.parse(TEST_PDF_FILE)
        self.assertEqual(len(lines), len(quotes))

    def test_pdf_ingestor_quotes_match_content(self):
        tmp = TEST_PDF_FILE.parent.joinpath('tmp.txt')
        subprocess.run(['pdftotext', '-layout',
                        TEST_PDF_FILE, tmp])
        with open(tmp) as f:
            line = f.readline()
            line = line.strip('\n\r').strip().split(' - ')
            pdf_quote_body = line[0].strip('"\'')
            pdf_quote_author = line[1]
        os.remove(tmp)

        quotes = Ingestor.parse(TEST_PDF_FILE)
        quote_obj_body = quotes[0].body.strip('"\'')
        quote_obj_author = quotes[0].author

        self.assertEqual(quote_obj_author, pdf_quote_author)
        self.assertEqual(quote_obj_body, pdf_quote_body)

    def test_pdf_ingestor_raises_error_for_wrong_filetype(self):
        self.assertRaises(Exception,
                          PDFIngestor.PDFIngestor.parse,
                          TEST_JSON_FILE)

    # DOCX Ingestor Tests #

    def test_docx_ingestor_returns_quote_objects(self):
        quotes = Ingestor.parse(TEST_DOCX_FILE)
        self.assertIsInstance(quotes[0], QuoteModel)

    def test_docx_ingestor_returns_all_quotes(self):
        doc = docx.Document(TEST_DOCX_FILE)
        lines = [line for line in doc.paragraphs if line.text != '']

        quotes = Ingestor.parse(TEST_DOCX_FILE)
        self.assertEqual(len(lines), len(quotes))

    def test_docx_ingestor_quotes_match_content(self):
        doc = docx.Document(TEST_DOCX_FILE)
        lines = [line.text for line in doc.paragraphs if line.text != '']
        doc_quote = lines[0].split(' - ')
        doc_quote_body = doc_quote[0].strip('"\'')
        doc_quote_author = doc_quote[1]

        quotes = Ingestor.parse(TEST_DOCX_FILE)
        quote_obj_body = quotes[0].body.strip('"\'')
        quote_obj_author = quotes[0].author

        self.assertEqual(quote_obj_author, doc_quote_author)
        self.assertEqual(quote_obj_body, doc_quote_body)

    def test_docx_ingestor_raises_error_for_wrong_filetype(self):
        self.assertRaises(Exception,
                          DocxIngestor.DocxIngestor.parse,
                          TEST_JSON_FILE)


if __name__ == '__main__':
    unittest.main()
